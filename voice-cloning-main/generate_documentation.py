"""
Script untuk generate dokumentasi Word dengan penjelasan kode MFCC
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Buat dokumen baru
doc = Document()

# ====================================================================
# COVER PAGE
# ====================================================================
title = doc.add_heading('DOKUMENTASI ANALISIS VOICE CLONING', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Ekstraksi dan Analisis Koefisien MFCC', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
info = doc.add_paragraph('Proyek: Voice Cloning MFCC Analysis')
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

info = doc.add_paragraph('Tanggal: 2024')
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ====================================================================
# DAFTAR ISI
# ====================================================================
doc.add_heading('Daftar Isi', 1)
doc.add_paragraph('1. Pendahuluan', style='List Number')
doc.add_paragraph('2. Perhitungan Nilai Statistik MFCC', style='List Number')
doc.add_paragraph('3. Analisis Perbandingan Fitur MFCC', style='List Number')
doc.add_paragraph('4. Visualisasi Data MFCC', style='List Number')
doc.add_paragraph('5. Interpretasi Hasil Analisis', style='List Number')
doc.add_paragraph('6. Penjelasan Kode Python', style='List Number')

doc.add_page_break()

# ====================================================================
# PENDAHULUAN
# ====================================================================
doc.add_heading('1. Pendahuluan', 1)
intro_text = """Mel-Frequency Cepstral Coefficients (MFCC) adalah teknik ekstraksi fitur yang digunakan untuk analisis karakteristik akustik suara. Penelitian ini membandingkan suara asli dengan hasil voice cloning untuk mengidentifikasi perbedaan karakteristik akustik.

Analisis dilakukan menggunakan 13 koefisien MFCC (N_MFCC = 13) yang diproses dari file audio dalam format WAV. Setiap pasangan rekaman (asli dan clone) dianalisis untuk mengetahui tingkat kesamaan karakteristik akustik."""
doc.add_paragraph(intro_text)

doc.add_page_break()

# ====================================================================
# SECTION 2: PERHITUNGAN STATISTIK
# ====================================================================
doc.add_heading('2. Perhitungan Nilai Statistik MFCC', 1)

doc.add_heading('2.1 Konsep Dasar', 2)
text = """Setelah ekstraksi MFCC menghasilkan matriks koefisien berukuran (13 × N) dimana 13 adalah jumlah koefisien dan N adalah jumlah frame, dilakukan perhitungan statistik:

a. Nilai Rata-rata (Mean): Menunjukkan kecenderungan nilai koefisien MFCC
b. Standar Deviasi (Std Dev): Menunjukkan tingkat variasi/perubahan nilai koefisien
c. Min dan Max: Menunjukkan range nilai koefisien
d. Median: Menunjukkan nilai tengah distribusi koefisien"""
doc.add_paragraph(text)

doc.add_heading('2.2 Hasil Statistik Contoh (asli1.wav)', 2)
text = """Berikut adalah statistik lengkap koefisien MFCC dari file asli1.wav:"""
doc.add_paragraph(text)

# Tambah tabel statistik
table = doc.add_table(rows=14, cols=6)
table.style = 'Light Grid Accent 1'
table.autofit = False

# Header
header_cells = table.rows[0].cells
headers = ['Koefisien', 'Mean', 'Std Dev', 'Min', 'Max', 'Median']
for i, header_text in enumerate(headers):
    header_cells[i].text = header_text
    header_cells[i].paragraphs[0].runs[0].font.bold = True

# Data statistik
stats_data = [
    ('MFCC-1', '-395.49', '67.61', '-571.61', '-246.11', '-393.51'),
    ('MFCC-2', '161.06', '53.63', '0.00', '257.37', '167.82'),
    ('MFCC-3', '43.23', '32.66', '-64.61', '122.91', '42.98'),
    ('MFCC-4', '22.38', '27.76', '-45.90', '102.23', '23.77'),
    ('MFCC-5', '13.09', '23.77', '-53.67', '78.76', '13.69'),
    ('MFCC-6', '-7.23', '18.00', '-59.33', '31.82', '-7.71'),
    ('MFCC-7', '4.70', '18.75', '-45.99', '77.02', '6.17'),
    ('MFCC-8', '9.63', '11.96', '-25.03', '50.06', '9.91'),
    ('MFCC-9', '-7.57', '14.35', '-46.10', '31.90', '-7.31'),
    ('MFCC-10', '-5.64', '13.29', '-47.92', '29.85', '-6.03'),
    ('MFCC-11', '8.31', '12.04', '-28.45', '41.83', '8.29'),
    ('MFCC-12', '4.34', '11.06', '-31.91', '34.78', '4.09'),
    ('MFCC-13', '-6.97', '8.84', '-29.08', '22.95', '-8.02'),
]

for i, (coef, mean, std, min_val, max_val, median) in enumerate(stats_data, 1):
    cells = table.rows[i].cells
    cells[0].text = coef
    cells[1].text = mean
    cells[2].text = std
    cells[3].text = min_val
    cells[4].text = max_val
    cells[5].text = median

explanation = """
Penjelasan Tabel:
• MFCC-1: Merupakan koefisien energi tertinggi dengan mean -395.49, menunjukkan energi dominan dalam spektrum mel
• MFCC-2: Mean 161.06, merupakan kontribusi frekuensi menengah ke tinggi
• MFCC-3-13: Merepresentasikan detail frekuensi lebih tinggi dengan kontribusi energi lebih kecil
• Std Dev yang tinggi pada MFCC-1 dan MFCC-2 menunjukkan variasi energi yang signifikan selama produksi suara
• Std Dev yang rendah pada MFCC-13 menunjukkan koefisien ini lebih stabil dalam karakterisasi suara"""
doc.add_paragraph(explanation)

doc.add_page_break()

# ====================================================================
# SECTION 3: ANALISIS PERBANDINGAN
# ====================================================================
doc.add_heading('3. Analisis Perbandingan Fitur MFCC', 1)

doc.add_heading('3.1 Ringkasan Hasil Perbandingan', 2)
text = """Analisis perbandingan dilakukan antara suara asli dan hasil voice cloning. Indikator kesamaan diukur dari:
a. Perbedaan rata-rata (Mean Difference): Semakin kecil semakin mirip
b. Perbedaan standar deviasi (Std Difference): Semakin kecil semakin mirip
c. Pola distribusi MFCC: Evaluasi visual dari heatmap MFCC"""
doc.add_paragraph(text)

# Tabel ringkasan
table2 = doc.add_table(rows=7, cols=5)
table2.style = 'Light Grid Accent 1'

header_cells2 = table2.rows[0].cells
headers2 = ['Pasangan', 'File Asli', 'File Clone', 'Rata-rata Selisih Mean', 'Rata-rata Selisih Std']
for i, h in enumerate(headers2):
    header_cells2[i].text = h
    header_cells2[i].paragraphs[0].runs[0].font.bold = True

data = [
    ('1', 'asli1.wav', 'clone1.wav', '9.62', '4.75'),
    ('2', 'asli2.wav', 'clone2.wav', '9.51', '3.04'),
    ('3', 'asli3.wav', 'clone3.wav', '8.59', '4.33'),
    ('4', 'asli4.wav', 'clone4.wav', '10.27', '2.58'),
    ('5', 'asli5.wav', 'clone5.wav', '10.22', '3.23'),
    ('6', 'asli_cowo.wav', 'clone_cowo.wav', '7.50', '2.17'),
]

for idx, (pair, asli, clone, mean_diff, std_diff) in enumerate(data, 1):
    cells = table2.rows[idx].cells
    cells[0].text = pair
    cells[1].text = asli
    cells[2].text = clone
    cells[3].text = mean_diff
    cells[4].text = std_diff

explanation2 = """
Interpretasi Hasil:
• Pasangan 6 (asli_cowo__clone_cowo) menunjukkan kesamaan tertinggi dengan selisih mean terendah (7.50)
• Pasangan 4 (asli4__clone4) menunjukkan selisih mean tertinggi (10.27) namun std diff rendah
• Semua pasangan memiliki perbedaan std yang lebih kecil dibanding perbedaan mean
• Nilai selisih yang lebih kecil menunjukkan kemiripan yang lebih tinggi antara suara asli dengan cloning"""
doc.add_paragraph(explanation2)

doc.add_page_break()

# ====================================================================
# SECTION 4: VISUALISASI
# ====================================================================
doc.add_heading('4. Visualisasi Data MFCC', 1)

doc.add_heading('4.1 Jenis-jenis Visualisasi', 2)
text = """Analisis visual dilakukan menggunakan beberapa jenis grafik:

a. Waveform: Menampilkan bentuk gelombang sinyal suara dalam domain waktu
b. Spectrogram: Menampilkan distribusi energi frekuensi terhadap waktu
c. MFCC Heatmap: Menampilkan intensitas koefisien MFCC dalam bentuk peta warna
d. Perbandingan Mean MFCC: Bar chart membandingkan nilai mean antar koefisien
e. Perbandingan Std MFCC: Bar chart membandingkan standar deviasi antar koefisien
f. Distribusi Mean: Histogram distribusi nilai mean koefisien
g. Variasi Frame Temporal: Line plot menunjukkan variasi antar frame"""
doc.add_paragraph(text)

doc.add_page_break()

# ====================================================================
# SECTION 5: PENJELASAN KODE
# ====================================================================
doc.add_heading('5. Penjelasan Kode Python', 1)

# Function 1: Load Audio
doc.add_heading('5.1 Fungsi: load_audio()', 2)
code_text = """def load_audio(file_path):
    \"\"\"Load audio file dengan error handling.\"\"\"
    try:
        y, sr = librosa.load(file_path, sr=None)
        duration = len(y) / sr
        return y, sr, duration
    except Exception as e:
        print(f"ERROR memuat {file_path}: {e}")
        return None, None, None"""
code_para = doc.add_paragraph(code_text, style='Normal')
for run in code_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

explanation = """
Penjelasan:
• Fungsi ini memuat file audio menggunakan librosa.load()
• Parameter sr=None berarti menggunakan sampling rate asli dari file
• y: array sinyal audio dalam bentuk time series
• sr: sampling rate (Hz) dari file audio
• duration: durasi audio dalam detik (dihitung dari total sampel / sampling rate)
• Error handling memastikan program tidak crash jika file tidak ditemukan"""
doc.add_paragraph(explanation)

doc.add_page_break()

# Function 2: Extract MFCC
doc.add_heading('5.2 Fungsi: extract_mfcc()', 2)
code_text2 = """def extract_mfcc(y, sr, n_mfcc=N_MFCC):
    \"\"\"Ekstraksi MFCC menggunakan librosa.\"\"\"
    try:
        mfcc = librosa.feature.mfcc(y=y, sr=sr, n_mfcc=n_mfcc)
        mfcc_mean = np.mean(mfcc, axis=1)
        mfcc_std = np.std(mfcc, axis=1)
        frame_var = np.var(mfcc, axis=0)
        return mfcc, mfcc_mean, mfcc_std, frame_var
    except Exception as e:
        print(f"ERROR ekstraksi MFCC: {e}")
        return None, None, None, None"""
code_para = doc.add_paragraph(code_text2, style='Normal')
for run in code_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

explanation = """
Penjelasan:
• librosa.feature.mfcc(): Ekstraksi 13 koefisien MFCC dari sinyal audio
• mfcc: Matriks berukuran (n_mfcc × num_frames) berisi koefisien MFCC
• axis=1: Menghitung statistik di sepanjang sumbu frame (waktu)
• mfcc_mean: Rata-rata setiap koefisien MFCC (nilai representatif)
• mfcc_std: Standar deviasi setiap koefisien MFCC (variabilitas)
• frame_var: Varians per frame menunjukkan dinamika temporal suara
• Proses ini mengkonversi informasi akustik kompleks menjadi 13 parameter kunci"""
doc.add_paragraph(explanation)

doc.add_page_break()

# Function 3: Plot Heatmap
doc.add_heading('5.3 Fungsi: plot_mfcc()', 2)
code_text3 = """def plot_mfcc(mfcc, sr, title, save_path):
    \"\"\"Plot MFCC heatmap.\"\"\"
    mfcc_db = librosa.amplitude_to_db(np.abs(mfcc), ref=np.max)
    
    plt.figure(figsize=(12, 5))
    librosa.display.specshow(mfcc_db, sr=sr, x_axis='time', y_axis='mel', cmap='coolwarm')
    plt.colorbar(format='%+2.0f dB', label='Magnitude (dB)')
    plt.title(title, fontsize=13, fontweight='bold')
    plt.xlabel("Waktu (detik)", fontsize=11)
    plt.ylabel("Koefisien MFCC", fontsize=11)
    plt.yticks(range(N_MFCC), [f'MFCC-{i+1}' for i in range(N_MFCC)], fontsize=9)
    plt.tight_layout()
    plt.savefig(save_path, dpi=DPI, bbox_inches='tight')
    plt.close()"""
code_para = doc.add_paragraph(code_text3, style='Normal')
for run in code_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

explanation = """
Penjelasan:
• librosa.amplitude_to_db(): Konversi amplitude ke skala desibel (dB) untuk visualisasi yang lebih baik
• specshow(): Menampilkan spektogram dengan colormap 'coolwarm' (biru=rendah, merah=tinggi)
• x_axis='time': Sumbu horizontal menunjukkan waktu dalam detik
• y_axis='mel': Sumbu vertikal menunjukkan skala mel (frekuensi perceptual)
• Colorbar: Menunjukkan skala intensitas dalam dB
• Heatmap ini memudahkan visualisasi pola karakteristik akustik suara dalam dimensi waktu-frekuensi"""
doc.add_paragraph(explanation)

doc.add_page_break()

# Function 4: Compare Means
doc.add_heading('5.4 Fungsi: plot_mean_comparison()', 2)
code_text4 = """def plot_mean_comparison(mean_asli, mean_clone, title, save_path):
    \"\"\"Plot perbandingan mean MFCC coefficients.\"\"\"
    x = np.arange(1, N_MFCC + 1)
    width = 0.35
    
    plt.figure(figsize=(12, 5))
    plt.bar(x - width/2, mean_asli, width, label='Suara Asli', color='steelblue', alpha=0.8)
    plt.bar(x + width/2, mean_clone, width, label='Voice Cloning', color='coral', alpha=0.8)
    
    plt.title(title, fontsize=13, fontweight='bold')
    plt.xlabel("Koefisien MFCC", fontsize=11)
    plt.ylabel("Nilai Mean", fontsize=11)
    plt.xticks(x, [f'MFCC-{i}' for i in x], fontsize=9)
    plt.legend(fontsize=10)
    plt.grid(True, alpha=0.3, axis='y')
    plt.tight_layout()
    plt.savefig(save_path, dpi=DPI, bbox_inches='tight')
    plt.close()"""
code_para = doc.add_paragraph(code_text4, style='Normal')
for run in code_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

explanation = """
Penjelasan:
• np.arange(): Membuat array posisi x-axis untuk 13 koefisien MFCC
• width=0.35: Lebar setiap bar chart untuk membandingkan dua kategori
• x - width/2 dan x + width/2: Menempatkan bar asli dan clone berdampingan
• color='steelblue' dan 'coral': Warna kontras untuk membedakan kategori
• grid(axis='y'): Garis grid horizontal untuk memudahkan membaca nilai
• Grafik ini memudahkan identifikasi koefisien mana yang berbeda signifikan antara asli dan clone"""
doc.add_paragraph(explanation)

doc.add_page_break()

# Function 5: Analyze Pair
doc.add_heading('5.5 Fungsi: analyze_pair()', 2)
code_snippet = """def analyze_pair(file_asli, file_clone, idx):
    \"\"\"Analisis pasangan file asli dan clone.\"\"\"
    # Load audio
    y_asli, sr_asli, dur_asli = load_audio(file_asli)
    y_clone, sr_clone, dur_clone = load_audio(file_clone)
    
    # Ekstraksi MFCC
    mfcc_asli, mean_asli, std_asli, ... = extract_mfcc(y_asli, sr_asli)
    mfcc_clone, mean_clone, std_clone, ... = extract_mfcc(y_clone, sr_clone)
    
    # Simpan CSV dan visualisasi untuk setiap pasangan"""
code_para = doc.add_paragraph(code_snippet, style='Normal')
for run in code_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

explanation = """
Penjelasan:
• Fungsi ini merupakan core function untuk analisis setiap pasangan asli-clone
• Workflow: Load → Ekstraksi MFCC → Perhitungan Statistik → Visualisasi
• CSV disimpan berisi 13 koefisien MFCC dengan mean, std, dan perbedaannya
• Menghasilkan 9 jenis visualisasi per pasangan:
  - 2 waveform (asli + clone)
  - 2 spectrogram (asli + clone)
  - 2 MFCC heatmap (asli + clone)
  - 3 grafik perbandingan (mean, std, distribusi)
• Semua output tersimpan dalam folder 'hasil' untuk dokumentasi"""
doc.add_paragraph(explanation)

doc.add_page_break()

# ====================================================================
# KESIMPULAN
# ====================================================================
doc.add_heading('6. Kesimpulan', 1)

kesimpulan = """Analisis MFCC terhadap 6 pasangan suara asli dan hasil voice cloning menunjukkan:

1. Karakteristik Akustik: Voice cloning menghasilkan karakteristik akustik yang mendekati suara asli, terutama pada standar deviasi koefisien MFCC yang menunjukkan perbedaan rata-rata < 5%.

2. Variabilitas Energi: Koefisien MFCC-1 dan MFCC-2 menunjukkan perbedaan tertinggi antara asli dan clone, mengindikasikan perbedaan dalam energi dan distribusi frekuensi menengah.

3. Stabilitas Koefisien Tinggi: Koefisien MFCC yang lebih tinggi (MFCC-10 hingga MFCC-13) menunjukkan stabilitas dan kesamaan yang lebih baik antara suara asli dan clone.

4. Indikator Forensik: Pola MFCC yang konsisten dan perbedaan mean dalam range 7-10 menunjukkan voice cloning berkualitas tinggi namun masih dapat dibedakan melalui analisis spektral mendalam.

5. Implikasi: Hasil ini menunjukkan bahwa MFCC dapat digunakan sebagai indikator dalam analisis forensik audio untuk mendeteksi voice cloning."""
doc.add_paragraph(kesimpulan)

# ====================================================================
# SAVE DOCUMENT
# ====================================================================
output_path = os.path.join('hasil_full', 'Dokumentasi_Voice_Cloning_MFCC_Analysis.docx')
doc.save(output_path)
print(f"✓ Dokumen Word berhasil dibuat: {output_path}")
