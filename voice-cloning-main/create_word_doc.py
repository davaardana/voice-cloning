"""
Script untuk membuat Word document dengan penjelasan MFCC dan thesis.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# Path ke folder hasil
HASIL_FOLDER = "hasil"
HASIL_FULL_FOLDER = "hasil_full"

def add_heading(doc, text, level=1):
    """Tambah heading dengan format yang baik."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False):
    """Tambah paragraf dengan format yang baik."""
    p = doc.add_paragraph(text)
    if bold:
        p.runs[0].bold = True
    return p

def add_image(doc, image_path, width=None):
    """Tambah gambar ke document."""
    if width is None:
        width = Inches(5.5)
    try:
        doc.add_picture(image_path, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Gambar tidak ditemukan: {image_path}]")

def create_word_document():
    """Buat Word document lengkap."""
    doc = Document()
    
    # Set font default
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ============================================
    # HALAMAN JUDUL
    # ============================================
    add_heading(doc, "LAPORAN ANALISIS VOICE CLONING", level=0)
    add_paragraph(doc, "Metode Ekstraksi MFCC (Mel-Frequency Cepstral Coefficients)")
    doc.add_paragraph()
    doc.add_paragraph("Disusun untuk Tugas Akhir / Thesis")
    doc.add_paragraph("Program Studi Teknik Informatika")
    doc.add_paragraph()
    doc.add_paragraph("=" * 50)
    
    # ============================================
    # DAFTAR ISI
    # ============================================
    add_heading(doc, "DAFTAR ISI", level=1)
    doc.add_paragraph("1. Pendahuluan")
    doc.add_paragraph("2. Teori MFCC")
    doc.add_paragraph("3. Penjelasan Kode Python")
    doc.add_paragraph("4. Hasil Analisis")
    doc.add_paragraph("5. Penjelasan Visualisasi")
    doc.add_paragraph("6. Kesimpulan")
    
    # ============================================
    # BAB 1: PENDAHULUAN
    # ============================================
    add_heading(doc, "1. PENDAHULUAN", level=1)
    add_paragraph(doc, "Voice cloning adalah teknologi untuk mereplikasi suara seseorang sehingga dapat menghasilkan audio yang terdengar seperti suara asli seseorang. Dalam penelitian ini, dilakukan analisis perbandingan antara suara asli dengan hasil voice cloning menggunakan metode MFCC (Mel-Frequency Cepstral Coefficients).", bold=True)
    
    add_paragraph(doc, "Tujuan analisis ini adalah:")
    doc.add_paragraph("• Mengidentifikasi perbedaan karakteristik antara suara asli dan suara hasil cloning")
    doc.add_paragraph("• Mengukur tingkat kemiripan menggunakan koefisien MFCC")
    doc.add_paragraph("• Menyediakan visualisasi untuk memahami proses ekstraksi fitur audio")
    
    # ============================================
    # BAB 2: TEORI MFCC
    # ============================================
    add_heading(doc, "2. TEORI MFCC", level=1)
    add_paragraph(doc, "MFCC (Mel-Frequency Cepstral Coefficients) adalah fitur yang paling umum digunakan dalam pemrosesan audio dan pengenalan suara. MFCC merepresentasikan karakteristik spektral suara berdasarkan skala Mel.", bold=True)
    
    add_heading(doc, "2.1 Tahapan Ekstraksi MFCC", level=2)
    
    add_paragraph(doc, "Proses ekstraksi MFCC terdiri dari langkah-langkah berikut:")
    
    # Tabel tahapan MFCC
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    
    steps = [
        ("1. Pre-emphasis", "Memperkuat frekuensi tinggi untuk menyeimbangkan spektrum. Formula: y[n] = x[n] - α*x[n-1], dengan α ≈ 0.97"),
        ("2. Framing", "Membagi sinyal audio menjadi frame-frame kecil (biasanya 20-40ms) dengan overlap 50%"),
        ("3. Windowing", "Mengaplikasi fungsi window (biasanya Hamming) untuk mengurangi diskontinuitas di ujung frame"),
        ("4. FFT (Fast Fourier Transform)", "Mengubah sinyal dari domain waktu ke frekuensi untuk mendapatkan spektrum"),
        ("5. Mel Filterbank", "Mengaplikasi filterbank dalam skala Mel untuk mensimulasikan pendengaran manusia"),
        ("6. Log", "Mengambil logaritma dari output Mel filterbank"),
        ("7. DCT (Discrete Cosine Transform)", "Mengurangi dimensi dan menghasilkan koefisien MFCC (biasanya 13-40 koefisien)")
    ]
    
    for i, (step, desc) in enumerate(steps):
        row = table.rows[i]
        row.cells[0].text = step
        row.cells[1].text = desc
    
    add_heading(doc, "2.2 Formula Matematika", level=2)
    
    add_paragraph(doc, "Berikut adalah formula matematika yang digunakan dalam ekstraksi MFCC:")
    
    # Formula
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Formula MFCC: ").bold = True
    p.add_run("c[n] = Σ(k=0 to K-1) log(S[k]) * cos(πn(k+0.5)/K)")
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Dimana: ").bold = True
    doc.add_paragraph("• c[n] = koefisien MFCC ke-n")
    doc.add_paragraph("• S[k] = output dari Mel filterbank")
    doc.add_paragraph("• K = jumlah filter dalam Mel filterbank")
    doc.add_paragraph("• n = nomor koefisien (1 hingga N_MFCC)")
    
    # Skala Mel
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Konversi Hz ke Mel: ").bold = True
    p.add_run("m = 2595 * log10(1 + f/700)")
    
    # ============================================
    # BAB 3: PENJELASAN KODE PYTHON
    # ============================================
    add_heading(doc, "3. PENJELASAN KODE PYTHON", level=1)
    add_paragraph(doc, "Berikut adalah penjelasan dari setiap script Python yang digunakan dalam analisis ini:", bold=True)
    
    add_heading(doc, "3.1 analisis.py - Ekstraksi MFCC", level=2)
    
    add_paragraph(doc, "Script ini berfungsi untuk mengekstraksi fitur MFCC dari semua file audio dalam folder dataset.")
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    params = [
        ("N_MFCC", "Jumlah koefisien MFCC yang diekstrak (default: 13)"),
        ("SR", "Sample rate audio (None = otomatis dari file)"),
        ("DATASET_FOLDER", "Folder tempat file audio berada"),
        ("OUTPUT_CSV", "File output untuk hasil ekstraksi MFCC"),
        ("SUMMARY_CSV", "File output untuk ringkasan statistik per kategori")
    ]
    
    for i, (param, desc) in enumerate(params):
        row = table.rows[i]
        row.cells[0].text = param
        row.cells[1].text = desc
    
    add_paragraph(doc, "Fungsi utama: extract_mfcc(audio_path, n_mfcc)")
    add_paragraph(doc, "Fungsi ini menggunakan library librosa untuk mengekstraksi MFCC dengan langkah-langkah:")
    doc.add_paragraph("1. Load audio menggunakan librosa.load()")
    doc.add_paragraph("2. Ekstraksi MFCC menggunakan librosa.feature.mfcc()")
    doc.add_paragraph("3. Hitung mean dan std dari setiap koefisien")
    doc.add_paragraph("4. Return nilai mean dan std sebagai feature vector")
    
    add_heading(doc, "3.2 analisis_mfcc.py - Analisis Pair-wise", level=2)
    
    add_paragraph(doc, "Script ini melakukan analisis perbandingan antara file audio asli dengan file cloning secara berpasangan.")
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    visualizations = [
        ("Waveform", "Visualisasi bentuk gelombang audio dalam domain waktu"),
        ("Spectrogram", "Visualisasi spektrogram frekuensi-waktu"),
        ("MFCC Heatmap", "Visualisasi koefisien MFCC setiap frame"),
        ("Perbandingan Mean/Std", "Grafik perbandingan nilai mean dan std antar koefisien MFCC")
    ]
    
    for i, (viz, desc) in enumerate(visualizations):
        row = table.rows[i]
        row.cells[0].text = viz
        row.cells[1].text = desc
    
    add_heading(doc, "3.3 analisis_full.py - Step-by-Step MFCC", level=2)
    
    add_paragraph(doc, "Script ini menampilkan tahapan demi tahapan proses ekstraksi MFCC secara detail untuk satu file audio.")
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    
    stages = [
        ("01_waveform_asli.png", "Bentuk gelombang audio dalam domain waktu"),
        ("02_mel_spectrogram.png", "Spektrogram dalam skala Mel"),
        ("03_mfcc_heatmap.png", "Heatmap koefisien MFCC per frame"),
        ("04_mean_std_mfcc.png", "Grafik mean dan std setiap koefisien MFCC"),
        ("05_temporal_variation.png", "Variasi nilai MFCC sepanjang waktu")
    ]
    
    for i, (file, desc) in enumerate(stages):
        row = table.rows[i]
        row.cells[0].text = file
        row.cells[1].text = desc
    
    # ============================================
    # BAB 4: HASIL ANALISIS
    # ============================================
    add_heading(doc, "4. HASIL ANALISIS", level=1)
    add_paragraph(doc, "Berikut adalah ringkasan hasil analisis perbandingan antara suara asli dengan voice cloning:", bold=True)
    
    # Tabel hasil perbandingan
    table = doc.add_table(rows=7, cols=5)
    table.style = 'Table Grid'
    
    # Header
    header_row = table.rows[0]
    header_row.cells[0].text = "Pasangan"
    header_row.cells[1].text = "File Asli"
    header_row.cells[2].text = "File Clone"
    header_row.cells[3].text = "Rata-rata Selisih Mean"
    header_row.cells[4].text = "Rata-rata Selisih Std"
    
    # Data dari ringkasan_pasangan_mfcc.csv
    data = [
        ("1", "asli1.wav", "clone1.wav", "9.62", "4.75"),
        ("2", "asli2.wav", "clone2.wav", "9.51", "3.04"),
        ("3", "asli3.wav", "clone3.wav", "8.59", "4.33"),
        ("4", "asli4.wav", "clone4.wav", "10.27", "2.58"),
        ("5", "asli5.wav", "clone5.wav", "10.22", "3.23"),
        ("6", "asli_cowo.wav", "clone_cowo.wav", "7.50", "2.17")
    ]
    
    for i, (pair, asli, clone, mean_diff, std_diff) in enumerate(data):
        row = table.rows[i+1]
        row.cells[0].text = pair
        row.cells[1].text = asli
        row.cells[2].text = clone
        row.cells[3].text = mean_diff
        row.cells[4].text = std_diff
    
    add_paragraph(doc, "Catatan: Nilai selisih yang lebih kecil menunjukkan kemiripan yang lebih tinggi antara suara asli dengan cloning.")
    
    # Tabel statistik MFCC
    add_heading(doc, "4.1 Statistik MFCC (Contoh: asli1.wav)", level=2)
    
    table = doc.add_table(rows=14, cols=4)
    table.style = 'Table Grid'
    
    # Header
    header_row = table.rows[0]
    header_row.cells[0].text = "Koefisien"
    header_row.cells[1].text = "Mean"
    header_row.cells[2].text = "Std Dev"
    header_row.cells[3].text = "Median"
    
    # Data dari mfcc_statistics.csv
    mfcc_data = [
        ("MFCC-1", "-395.49", "67.61", "-393.51"),
        ("MFCC-2", "161.06", "53.63", "167.82"),
        ("MFCC-3", "43.23", "32.66", "42.98"),
        ("MFCC-4", "22.38", "27.76", "23.77"),
        ("MFCC-5", "13.09", "23.77", "13.69"),
        ("MFCC-6", "-7.23", "18.00", "-7.71"),
        ("MFCC-7", "4.70", "18.75", "6.17"),
        ("MFCC-8", "9.63", "11.96", "9.91"),
        ("MFCC-9", "-7.57", "14.35", "-7.31"),
        ("MFCC-10", "-5.64", "13.29", "-6.03"),
        ("MFCC-11", "8.31", "12.04", "8.29"),
        ("MFCC-12", "4.34", "11.06", "4.09"),
        ("MFCC-13", "-6.97", "8.84", "-8.02")
    ]
    
    for i, (coef, mean, std, median) in enumerate(mfcc_data):
        row = table.rows[i+1]
        row.cells[0].text = coef
        row.cells[1].text = mean
        row.cells[2].text = std
        row.cells[3].text = median
    
    # ============================================
    # BAB 5: PENJELASAN VISUALISASI
    # ============================================
    add_heading(doc, "5. PENJELASAN VISUALISASI", level=1)
    add_paragraph(doc, "Berikut adalah penjelasan untuk setiap gambar yang dihasilkan:", bold=True)
    
    add_heading(doc, "5.1 Visualisasi di Folder hasil_full/", level=2)
    
    # Penjelasan gambar hasil_full
    full_viz = [
        ("01_waveform_asli.png", "Menampilkan bentuk gelombang audio dalam domain waktu. Sumbu X menunjukkan waktu (detik) dan sumbu Y menunjukkan amplitude. Dari waveform dapat dilihat karakteristik umum suara seperti durasi, volume, dan pola perubahan amplitude."),
        ("02_mel_spectrogram.png", "Menampilkan spektrogram dalam skala Mel. Skala Mel lebih akurat dalam merepresentasikan persepsi pendengaran manusia. Warna menunjukkan intensitas energi pada frekuensi tertentu."),
        ("03_mfcc_heatmap.png", "Menampilkan heatmap koefisien MFCC. Sumbu X adalah nomor frame dan sumbu Y adalah nomor koefisien MFCC. Warna menunjukkan nilai koefisien. Koefisien pertama (MFCC-1) biasanya menunjukkan rata-rata energi."),
        ("04_mean_std_mfcc.png", "Menampilkan grafik mean dan standar deviasi untuk setiap koefisien MFCC. Bar menunjukkan nilai mean dengan error bar menunjukkan standar deviasi. Ini membantu melihat distribusi setiap koefisien."),
        ("05_temporal_variation.png", "Menampilkan variasi nilai MFCC sepanjang waktu untuk setiap koefisien. Garis menunjukkan bagaimana nilai koefisien berubah dari frame ke frame.")
    ]
    
    for i, (file, desc) in enumerate(full_viz):
        add_paragraph(doc, f"{i+1}. {file}", bold=True)
        add_paragraph(doc, desc)
    
    add_heading(doc, "5.2 Visualisasi di Folder hasil/", level=2)
    
    add_paragraph(doc, "Folder hasil/ berisi visualisasi perbandingan untuk setiap pasangan asli-clone. Untuk setiap pasangan dihasilkan 9 gambar:")
    
    table = doc.add_table(rows=10, cols=2)
    table.style = 'Table Grid'
    
    pair_viz = [
        ("waveform_asli_*.png", "Bentuk gelombang suara asli"),
        ("waveform_clone_*.png", "Bentuk gelombang suara cloning"),
        ("spectrogram_asli_*.png", "Spektrogram suara asli"),
        ("spectrogram_clone_*.png", "Spektrogram suara cloning"),
        ("mfcc_asli_*.png", "Heatmap MFCC suara asli"),
        ("mfcc_clone_*.png", "Heatmap MFCC suara cloning"),
        ("perbandingan_mean_mfcc_*.png", "Grafik perbandingan mean MFCC asli vs clone"),
        ("perbandingan_std_mfcc_*.png", "Grafik perbandingan std MFCC asli vs clone"),
        ("distribusi_mean_mfcc_*.png", "Distribusi nilai mean MFCC")
    ]
    
    for i, (file, desc) in enumerate(pair_viz):
        row = table.rows[i]
        row.cells[0].text = file
        row.cells[1].text = desc
    
    # ============================================
    # BAB 6: KESIMPULAN
    # ============================================
    add_heading(doc, "6. KESIMPULAN", level=1)
    add_paragraph(doc, "Berdasarkan analisis yang telah dilakukan, dapat ditarik kesimpulan:", bold=True)
    
    doc.add_paragraph("1. Metode MFCC efektif untuk menganalisis karakteristik suara dan membedakan antara suara asli dengan hasil voice cloning.")
    doc.add_paragraph("2. Koefisien MFCC-1 (yang merepresentasikan energi rata-rata) menunjukkan perbedaan yang signifikan antara suara asli dan cloning.")
    doc.add_paragraph("3. Visualisasi heatmap MFCC membantu dalam melihat pola distribusi koefisien sepanjang waktu.")
    doc.add_paragraph("4. Perbandingan mean dan std antar koefisien memberikan informasi kuantitatif tentang tingkat kemiripan.")
    doc.add_paragraph("5.声音 cloning dengan kualitas tinggi akan memiliki nilai selisih mean dan std yang lebih kecil.")
    
    # ============================================
    # LAMPIRAN
    # ============================================
    add_heading(doc, "LAMPIRAN", level=1)
    add_paragraph(doc, "Daftar file yang dihasilkan:", bold=True)
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    row = table.rows[0]
    row.cells[0].text = "Folder"
    row.cells[1].text = "Jumlah File"
    
    row = table.rows[1]
    row.cells[0].text = "hasil/"
    row.cells[1].text = "60 file (6 pasangan × 10 visualisasi)"
    
    # Simpan document
    output_path = "laporan_voice_cloning.docx"
    doc.save(output_path)
    print(f"Word document berhasil dibuat: {output_path}")

if __name__ == "__main__":
    create_word_document()